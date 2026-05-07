import zipfile
import traceback
import threading
from pathlib import Path
from typing import List

from fastapi import APIRouter, File, UploadFile, HTTPException, BackgroundTasks, Query
from fastapi.responses import FileResponse

from app.utils import (
    unique_path,
    cleanup_files,
    validate_ext,
    save_upload_file,
    create_job,
    update_job,
    attach_job_files,
    complete_job,
    fail_job,
)

router = APIRouter(prefix="/convert", tags=["PDF"])


# ── PDF → Word ────────────────────────────────────────────────────────────

@router.post("/pdf-to-word")
async def convert_pdf_to_word(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    ocr: bool = Query(False, description="Ativa OCR quando o PDF não possui texto extraível"),
    ocr_lang: str = Query("por", description="Idioma do OCR (ex.: por, eng, por+eng)"),
    ocr_dpi: int = Query(200, description="Resolução (DPI) usada para OCR"),
):
    validate_ext(file.filename, {".pdf"})
    input_path = unique_path(".pdf")
    output_path = input_path.with_suffix(".docx")
    try:
        await save_upload_file(file, input_path)

        def docx_has_meaningful_text(path: Path) -> bool:
            try:
                from docx import Document
            except Exception:
                return False

            try:
                doc = Document(str(path))
            except Exception:
                return False

            text_parts: list[str] = []

            for p in doc.paragraphs:
                if p.text:
                    text_parts.append(p.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_parts.append(cell.text)

            combined = "\n".join(text_parts).strip()
            return len(combined) > 0

        def build_docx_from_pdf_text(pdf_path: Path, docx_path: Path) -> None:
            import fitz
            from docx import Document

            pdf = fitz.open(str(pdf_path))
            doc = Document()

            for i in range(len(pdf)):
                page = pdf[i]
                words = page.get_text("words") or []
                lines: dict[tuple[int, int], dict] = {}
                for w in words:
                    x0, y0, x1, y1, text, block_no, line_no, word_no = w
                    text = (text or "").strip()
                    if not text:
                        continue
                    key = (int(block_no), int(line_no))
                    item = lines.get(key)
                    if not item:
                        item = {"y0": float(y0), "x0": float(x0), "parts": []}
                        lines[key] = item
                    else:
                        if float(y0) < item["y0"]:
                            item["y0"] = float(y0)
                        if float(x0) < item["x0"]:
                            item["x0"] = float(x0)
                    item["parts"].append((float(x0), text))

                ordered = sorted(
                    lines.values(),
                    key=lambda it: (it["y0"], it["x0"]),
                )

                rendered_lines: list[str] = []
                for it in ordered:
                    parts = sorted(it["parts"], key=lambda p: p[0])
                    line_text = " ".join(t for _, t in parts).strip()
                    if line_text:
                        rendered_lines.append(line_text)

                normalized_lines: list[str] = []
                for line_text in rendered_lines:
                    if (
                        normalized_lines
                        and normalized_lines[-1].endswith("-")
                        and line_text[:1].islower()
                    ):
                        normalized_lines[-1] = normalized_lines[-1][:-1] + line_text
                    else:
                        normalized_lines.append(line_text)

                for line_text in normalized_lines:
                    if line_text:
                        doc.add_paragraph(line_text)
                if i < len(pdf) - 1:
                    doc.add_page_break()

            pdf.close()
            doc.save(str(docx_path))

        def build_docx_from_pdf_ocr(pdf_path: Path, docx_path: Path) -> None:
            try:
                import pytesseract
            except Exception as exc:
                raise HTTPException(
                    500,
                    f"OCR não está disponível no servidor (pytesseract). Erro: {exc}",
                )

            try:
                import fitz
            except Exception as exc:
                raise HTTPException(500, f"Falha ao carregar PyMuPDF para OCR: {exc}")

            try:
                from PIL import Image
            except Exception as exc:
                raise HTTPException(500, f"Falha ao carregar Pillow para OCR: {exc}")

            from docx import Document

            pdf = fitz.open(str(pdf_path))
            doc = Document()
            scale = max(1.0, ocr_dpi / 72)
            matrix = fitz.Matrix(scale, scale)

            for i in range(len(pdf)):
                page = pdf[i]
                pix = page.get_pixmap(matrix=matrix, alpha=False)
                if pix.n == 1:
                    mode = "L"
                elif pix.n == 3:
                    mode = "RGB"
                elif pix.n == 4:
                    mode = "RGBA"
                else:
                    mode = "RGB"

                img = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
                if img.mode != "RGB":
                    img = img.convert("RGB")

                text = (pytesseract.image_to_string(img, lang=ocr_lang) or "").strip()
                if text:
                    for line in text.splitlines():
                        line = line.strip()
                        if line:
                            doc.add_paragraph(line)
                if i < len(pdf) - 1:
                    doc.add_page_break()

            pdf.close()
            doc.save(str(docx_path))

        pdf2docx_error: Exception | None = None
        try:
            from pdf2docx import Converter as PdfToDocxConverter

            cv = PdfToDocxConverter(str(input_path))
            cv.convert(str(output_path), start=0, end=None)
            cv.close()
        except Exception as err:
            pdf2docx_error = err

        if output_path.exists() and not docx_has_meaningful_text(output_path):
            cleanup_files(output_path)

        if (not output_path.exists()) or (not docx_has_meaningful_text(output_path)):
            try:
                build_docx_from_pdf_text(input_path, output_path)
            except Exception:
                cleanup_files(output_path)

        if output_path.exists() and not docx_has_meaningful_text(output_path) and ocr:
            cleanup_files(output_path)
            build_docx_from_pdf_ocr(input_path, output_path)

        if output_path.exists() and not docx_has_meaningful_text(output_path):
            cleanup_files(output_path)
            if pdf2docx_error:
                raise HTTPException(
                    422,
                    "PDF parece ser escaneado ou não possui texto extraível; use o parâmetro ocr=true para tentar OCR.",
                )
            raise HTTPException(
                422,
                "PDF não possui texto extraível; use o parâmetro ocr=true para tentar OCR.",
            )

        if not output_path.exists():
            # Try to find the output file with a different name
            for f in input_path.parent.iterdir():
                if f.stem.startswith(input_path.stem.split("_")[0]) and f.suffix == ".docx":
                    output_path = f
                    break
            else:
                raise HTTPException(500, "DOCX output not found after conversion.")

        background_tasks.add_task(cleanup_files, input_path, output_path)
        return FileResponse(
            path=str(output_path),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{Path(file.filename).stem}.docx",
        )
    except HTTPException:
        cleanup_files(input_path, output_path)
        raise
    except Exception as exc:
        cleanup_files(input_path, output_path)
        error_detail = f"Conversion failed: {exc}\n{traceback.format_exc()}"
        print(f"[pdf-to-word] ERROR: {error_detail}")
        raise HTTPException(500, detail=f"Conversion failed: {exc}")


@router.post("/pdf-to-word-job")
async def convert_pdf_to_word_job(
    file: UploadFile = File(...),
    engine: str = Query("auto", description="auto, pdf2docx, text, ocr"),
    ocr: bool = Query(False, description="Ativa OCR quando o PDF não possui texto extraível"),
    ocr_lang: str = Query("por", description="Idioma do OCR (ex.: por, eng, por+eng)"),
    ocr_dpi: int = Query(200, description="Resolução (DPI) usada para OCR"),
    start_page: int = Query(1, ge=1, description="Página inicial (1-based)"),
    end_page: int | None = Query(None, ge=1, description="Página final (1-based, inclusiva)"),
):
    validate_ext(file.filename, {".pdf"})
    input_path = unique_path(".pdf")
    output_path = input_path.with_suffix(".docx")

    await save_upload_file(file, input_path)

    job_id = create_job()
    attach_job_files(job_id, input_path=input_path)
    update_job(job_id, progress=0, status="running")

    download_name = f"{Path(file.filename).stem}.docx"
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def progress_cb(pct: int) -> None:
        update_job(job_id, progress=pct)

    def worker() -> None:
        out_path = output_path
        progress_stop = threading.Event()
        try:
            progress_cb(5)
            update_job(job_id, message="Analisando PDF…")

            import fitz

            pdf = fitz.open(str(input_path))
            total_pages = len(pdf) or 1
            pdf.close()

            sp = int(start_page)
            ep = int(end_page) if end_page is not None else int(total_pages)
            if sp < 1 or ep < 1 or sp > ep:
                raise HTTPException(400, "Intervalo de páginas inválido.")
            if sp > total_pages:
                raise HTTPException(400, f"start_page fora do limite. Total: {total_pages}.")
            if ep > total_pages:
                ep = int(total_pages)

            selected_pages = int(ep - sp + 1)

            eng = (engine or "auto").strip().lower()
            if eng not in ("auto", "pdf2docx", "text", "ocr"):
                raise HTTPException(400, "engine inválido. Use: auto, pdf2docx, text, ocr.")

            if eng == "auto":
                eng = "text" if selected_pages > 120 else "pdf2docx"

            if eng == "ocr":
                ocr = True

            def docx_has_meaningful_text(path: Path) -> bool:
                try:
                    from docx import Document
                except Exception:
                    return False

                try:
                    doc = Document(str(path))
                except Exception:
                    return False

                text_parts: list[str] = []

                for p in doc.paragraphs:
                    if p.text:
                        text_parts.append(p.text)

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text:
                                text_parts.append(cell.text)

                combined = "\n".join(text_parts).strip()
                return len(combined) > 0

            def build_docx_from_pdf_text(pdf_path: Path, docx_path: Path) -> None:
                import fitz
                from docx import Document

                pdf = fitz.open(str(pdf_path))
                doc = Document()
                total = max(1, selected_pages)

                for idx, page_i in enumerate(range(sp - 1, ep)):
                    page = pdf[page_i]
                    words = page.get_text("words") or []
                    lines: dict[tuple[int, int], dict] = {}
                    for w in words:
                        x0, y0, x1, y1, text, block_no, line_no, word_no = w
                        text = (text or "").strip()
                        if not text:
                            continue
                        key = (int(block_no), int(line_no))
                        item = lines.get(key)
                        if not item:
                            item = {"y0": float(y0), "x0": float(x0), "parts": []}
                            lines[key] = item
                        else:
                            if float(y0) < item["y0"]:
                                item["y0"] = float(y0)
                            if float(x0) < item["x0"]:
                                item["x0"] = float(x0)
                        item["parts"].append((float(x0), text))

                    ordered = sorted(
                        lines.values(),
                        key=lambda it: (it["y0"], it["x0"]),
                    )

                    rendered_lines: list[str] = []
                    for it in ordered:
                        parts = sorted(it["parts"], key=lambda p: p[0])
                        line_text = " ".join(t for _, t in parts).strip()
                        if line_text:
                            rendered_lines.append(line_text)

                    normalized_lines: list[str] = []
                    for line_text in rendered_lines:
                        if (
                            normalized_lines
                            and normalized_lines[-1].endswith("-")
                            and line_text[:1].islower()
                        ):
                            normalized_lines[-1] = normalized_lines[-1][:-1] + line_text
                        else:
                            normalized_lines.append(line_text)

                    for line_text in normalized_lines:
                        if line_text:
                            doc.add_paragraph(line_text)
                    if page_i < ep - 1:
                        doc.add_page_break()

                    progress_cb(35 + int(((idx + 1) / total) * 60))

                pdf.close()
                doc.save(str(docx_path))

            def build_docx_from_pdf_ocr(pdf_path: Path, docx_path: Path) -> None:
                try:
                    import pytesseract
                except Exception as exc:
                    raise HTTPException(
                        500,
                        f"OCR não está disponível no servidor (pytesseract). Erro: {exc}",
                    )

                try:
                    import fitz
                except Exception as exc:
                    raise HTTPException(500, f"Falha ao carregar PyMuPDF para OCR: {exc}")

                try:
                    from PIL import Image
                except Exception as exc:
                    raise HTTPException(500, f"Falha ao carregar Pillow para OCR: {exc}")

                from docx import Document

                pdf = fitz.open(str(pdf_path))
                doc = Document()
                total = max(1, selected_pages)
                scale = max(1.0, ocr_dpi / 72)
                matrix = fitz.Matrix(scale, scale)

                for idx, page_i in enumerate(range(sp - 1, ep)):
                    page = pdf[page_i]
                    pix = page.get_pixmap(matrix=matrix, alpha=False)
                    if pix.n == 1:
                        mode = "L"
                    elif pix.n == 3:
                        mode = "RGB"
                    elif pix.n == 4:
                        mode = "RGBA"
                    else:
                        mode = "RGB"

                    img = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
                    if img.mode != "RGB":
                        img = img.convert("RGB")

                    text = (pytesseract.image_to_string(img, lang=ocr_lang) or "").strip()
                    if text:
                        for line in text.splitlines():
                            line = line.strip()
                            if line:
                                doc.add_paragraph(line)
                    if page_i < ep - 1:
                        doc.add_page_break()

                    progress_cb(35 + int(((idx + 1) / total) * 60))

                pdf.close()
                doc.save(str(docx_path))

            pdf2docx_error: Exception | None = None
            if eng == "pdf2docx":
                update_job(job_id, message="Convertendo (pdf2docx)…")

                def tick_progress() -> None:
                    p = 10
                    while not progress_stop.is_set():
                        p = min(34, p + 1)
                        update_job(job_id, progress=p)
                        progress_stop.wait(2.0)

                threading.Thread(target=tick_progress, daemon=True).start()

                try:
                    from pdf2docx import Converter as PdfToDocxConverter

                    progress_cb(10)
                    cv = PdfToDocxConverter(str(input_path))
                    cv.convert(str(out_path), start=sp - 1, end=ep)
                    cv.close()
                    progress_cb(35)
                except Exception as err:
                    pdf2docx_error = err
                    cleanup_files(out_path)
                finally:
                    progress_stop.set()
            else:
                progress_cb(35)

            if out_path.exists() and not docx_has_meaningful_text(out_path):
                cleanup_files(out_path)

            if (not out_path.exists()) or (not docx_has_meaningful_text(out_path)):
                try:
                    update_job(job_id, message="Extraindo texto…")
                    build_docx_from_pdf_text(input_path, out_path)
                except Exception:
                    cleanup_files(out_path)

            if out_path.exists() and not docx_has_meaningful_text(out_path) and ocr:
                cleanup_files(out_path)
                update_job(job_id, message="Executando OCR…")
                build_docx_from_pdf_ocr(input_path, out_path)

            if out_path.exists() and not docx_has_meaningful_text(out_path):
                cleanup_files(out_path)
                if pdf2docx_error:
                    raise HTTPException(
                        422,
                        "PDF parece ser escaneado ou não possui texto extraível; use o parâmetro ocr=true para tentar OCR.",
                    )
                raise HTTPException(
                    422,
                    "PDF não possui texto extraível; use o parâmetro ocr=true para tentar OCR.",
                )

            if not out_path.exists():
                for f in input_path.parent.iterdir():
                    if (
                        f.stem.startswith(input_path.stem.split("_")[0])
                        and f.suffix == ".docx"
                    ):
                        out_path = f
                        break
                else:
                    raise HTTPException(500, "DOCX output not found after conversion.")

            update_job(job_id, message="Finalizando…")
            complete_job(
                job_id,
                output_path=out_path,
                download_name=download_name,
                media_type=media_type,
            )
        except Exception as exc:
            progress_stop.set()
            fail_job(job_id, message=str(exc))
            cleanup_files(input_path, out_path)

    threading.Thread(target=worker, daemon=True).start()

    return {
        "job_id": job_id,
        "status_url": f"/jobs/{job_id}",
        "download_url": f"/jobs/{job_id}/download",
    }


# ── PDF Merge ──────────────────────────────────────────────────────────────


@router.post("/merge-pdfs")
async def merge_pdfs(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...),
):
    if len(files) < 2:
        raise HTTPException(400, detail="Envie pelo menos 2 arquivos PDF para juntar.")

    input_paths: List[Path] = []
    output_path = unique_path(".pdf")
    try:
        for f in files:
            validate_ext(f.filename, {".pdf"})
            p = unique_path(".pdf")
            await save_upload_file(f, p)
            input_paths.append(p)

        import fitz

        out_doc = fitz.open()
        for p in input_paths:
            doc = fitz.open(str(p))
            out_doc.insert_pdf(doc)
            doc.close()

        out_doc.save(str(output_path))
        out_doc.close()

        background_tasks.add_task(cleanup_files, *input_paths, output_path)
        return FileResponse(
            path=str(output_path),
            media_type="application/pdf",
            filename="merged.pdf",
        )
    except HTTPException:
        cleanup_files(*input_paths, output_path)
        raise
    except Exception as exc:
        cleanup_files(*input_paths, output_path)
        raise HTTPException(500, detail=f"Conversion failed: {exc}")


# ── PDF → Image ───────────────────────────────────────────────────────────


@router.post("/pdf-to-image")
async def convert_pdf_to_image(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    fmt: str = Query("png", description="Output image format: png or jpg"),
    dpi: int = Query(150, description="Image resolution in DPI"),
):
    validate_ext(file.filename, {".pdf"})
    if fmt.lower() not in ("png", "jpg", "jpeg"):
        raise HTTPException(400, "Format must be 'png' or 'jpg'.")

    output_ext = ".png" if fmt.lower() == "png" else ".jpg"
    input_path = unique_path(".pdf")

    try:
        await save_upload_file(file, input_path)

        import fitz  # PyMuPDF
        doc = fitz.open(input_path)
        output_paths: List[Path] = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            matrix = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=matrix)
            page_path = unique_path(f"_page{page_num + 1}{output_ext}")
            pix.save(str(page_path))
            output_paths.append(page_path)

        doc.close()

        if len(output_paths) == 1:
            p = output_paths[0]
            background_tasks.add_task(cleanup_files, input_path, p)
            return FileResponse(
                path=str(p),
                media_type=f"image/{fmt.lower().replace('jpg', 'jpeg')}",
                filename=f"{Path(file.filename).stem}{output_ext}",
            )

        zip_path = unique_path(".zip")
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for p in output_paths:
                zf.write(p, p.name)

        all_paths = [input_path, zip_path] + output_paths
        background_tasks.add_task(cleanup_files, *all_paths)

        return FileResponse(
            path=str(zip_path),
            media_type="application/zip",
            filename=f"{Path(file.filename).stem}_pages.zip",
        )

    except HTTPException:
        cleanup_files(input_path)
        raise
    except Exception as exc:
        cleanup_files(input_path)
        raise HTTPException(500, detail=f"Conversion failed: {exc}")


# ── Image → PDF ───────────────────────────────────────────────────────────

from PIL import Image


@router.post("/image-to-pdf")
async def convert_image_to_pdf(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
):
    validate_ext(file.filename, {".jpg", ".jpeg", ".png"})
    ext = Path(file.filename).suffix.lower()
    input_path = unique_path(ext)
    output_path = input_path.with_suffix(".pdf")
    try:
        await save_upload_file(file, input_path)

        image = Image.open(input_path)
        if image.mode in ("RGBA", "P"):
            image = image.convert("RGB")
        image.save(output_path, "PDF", resolution=100.0)

        background_tasks.add_task(cleanup_files, input_path, output_path)
        return FileResponse(
            path=str(output_path),
            media_type="application/pdf",
            filename=f"{Path(file.filename).stem}.pdf",
        )
    except HTTPException:
        cleanup_files(input_path, output_path)
        raise
    except Exception as exc:
        cleanup_files(input_path, output_path)
        raise HTTPException(500, detail=f"Conversion failed: {exc}")
